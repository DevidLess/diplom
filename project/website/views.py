import operator
from io import BytesIO
from urllib.parse import quote

from absl.flags import ValidationError
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.http import HttpResponse, FileResponse
from django.shortcuts import redirect, render
from django.urls import reverse_lazy, reverse
from django.views.generic import (
    TemplateView,
    CreateView,
    ListView,
    DeleteView,
    DetailView,
)
from django.views.generic.edit import FormView, UpdateView
from django_filters.views import FilterView
from docx.enum.text import WD_ALIGN_PARAGRAPH
from python_docx_replace import docx_replace

from website.access_control import *
from website.filters import *
from website.models import *
from website.forms import *

from django.utils.html import escape
from docx import Document
from docx.shared import Inches

# **********************************************************************************************
# Проверка прав доступа к данным
# **********************************************************************************************


def get_court_order(request, pk):
    model = CourtOrder.objects.get(id=pk)
    if not (model.defendant.id == request.user.id or model.claimant.id == request.user.id):
        raise ValidationError("Access denied")
    return model


def get_template(request, pk):
    model = Template.objects.get(id=pk)
    # print(model.user)
    # if model.user.id != request.user.id:
    #     raise ValidationError('Access denied')
    return model


def get_cancellation_request(request, pk):
    model = CancellationRequest.objects.get(id=pk)
    # print(model.user)
    if model.user.id != request.user.id:
        raise ValidationError("Access denied")
    return model


# **********************************************************************************************
# Заполнение шаблона
# **********************************************************************************************


@login_required
def docx_fill_template(request, template_pk, court_order_pk=None):
    template = get_template(request, template_pk)
    court_order = None
    try:
        court_order = get_court_order(request, court_order_pk)
    except:
        pass
    user = request.user
    dictionary = {}

    if user:
        dictionary = {
            "пользователь": user.title_long(),
            "пользователь_кратко": user.title_short(),
            "пользователь_адрес": user.address,
            "пользователь_телефон": user.phone,
            "текущая_дата": datetime.now().__format__("%d.%m.%Y"),
        }

    if court_order:
        dictionary.update(
            {
                "ответчик": court_order.defendant.title_long() if court_order.defendant else "",
                "ответчик_кратко": court_order.defendant.title_short()
                if court_order.defendant
                else "",
                "ответчик_адрес": court_order.defendant.address if court_order.defendant else "",
                "ответчик_телефон": court_order.defendant.phone if court_order.defendant else "",
                "истец": court_order.claimant.title_long() if court_order.claimant else "",
                "истец_кратко": court_order.claimant.title_short() if court_order.claimant else "",
                "истец_адрес": court_order.claimant.address if court_order.claimant else "",
                "истец_телефон": court_order.claimant.phone if court_order.claimant else "",
                "судебный_участок": court_order.court_district.title
                if court_order.court_district
                else "",
                "судебный_участок_адрес": court_order.court_district.address
                if court_order.court_district
                else "",
                "судебный_участок_телефон": court_order.court_district.phone
                if court_order.court_district
                else "",
                "судебный_приказ_дата": court_order.date.__format__("%d.%m.%Y"),
                "судебный_приказ_номер": court_order.number,
            }
        )

    document = Document(template.file)
    docx_replace(document, **dictionary)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # response['Content-Disposition'] = 'attachment; filename=template_' + template.id.__str__() + '.docx'
    response["Content-Disposition"] = "attachment; filename*=utf-8''{}".format(
        quote(template.title + ".docx")
    )
    document.save(response)

    return response


# **********************************************************************************************
# Генерация документов
# **********************************************************************************************


@login_required
def docx_cancellation_request(request, pk):
    cancellation_request = get_cancellation_request(request, pk)
    document = Document()

    p = document.add_paragraph(cancellation_request.court_order.court_district.title)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph(cancellation_request.court_order.court_district.address)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph(cancellation_request.court_order.court_district.phone)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph("Ответчик: " + cancellation_request.court_order.defendant.__str__())
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph("")

    p = document.add_heading("Заявление об отмене судебного приказа", 1)
    # p = document.add_paragraph("Заявление об отмене судебного приказа")
    p.style.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph("")

    p = document.add_paragraph("")
    p = document.add_paragraph(cancellation_request.content)
    # p.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    p = document.add_paragraph("")
    # p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = document.add_paragraph(cancellation_request.court_order.defendant.__str__())
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = document.add_paragraph(cancellation_request.registration_date.__format__("%d.%m.%Y"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = (
        "attachment; filename=Request_" + cancellation_request.id.__str__() + ".docx"
    )
    # response['Content-Disposition'] = "attachment; filename*=utf-8''{}".format(quote('Пример.docx'))
    document.save(response)

    return response


# **********************************************************************************************
# Сайт
# **********************************************************************************************
class HomePageView(TemplateView):
    title = "Главная страница"
    template_name = "site/home.html"


class AboutPageView(TemplateView):
    title = "О системе"
    template_name = "site/about.html"


# **********************************************************************************************
# Профиль пользователя
# **********************************************************************************************
class UserProfileDetailView(LoginRequiredMixin, TemplateView):
    title = "Профиль"
    template_name = "user_profile/detail.html"

    def get_context_data(self, *args, **kwargs):
        # Получение базового контекста
        context = super(UserProfileDetailView, self).get_context_data(*args, **kwargs)
        
        # Текущий пользователь
        user = self.request.user
        context["model"] = user

        # Фильтруем судебные приказы по пользователю
        court_orders = CourtOrder.objects.filter(claimant=user) | CourtOrder.objects.filter(defendant=user)
        context["court_orders"] = court_orders.distinct()  # Убираем дубли

        # Фильтруем заявления об отмене судебных приказов по пользователю
        cancellation_requests = CancellationRequest.objects.filter(user=user)
        context["cancellation_requests"] = cancellation_requests

        # Передача заголовка
        context["view"] = {"title": self.title}

        return context


# class UserProfileUpdateView(LoginRequiredMixin, FormView):
#     title = "Редактирование профиля"
#     template_name = "user_profile/update.html"
#     form_class = UserProfileForm
#     success_url = "/user_profile/"

#     def get_form(self, form_class=None):
#         return UserProfileForm(instance=self.request.user, **self.get_form_kwargs())

# class UserProfileUpdateView(LoginRequiredMixin, FormView):
#     title = "Редактирование профиля"
#     template_name = "user_profile/update.html"
#     form_class = UserProfileForm
#     success_url = "/user_profile/"

#     def get_form(self, form_class=None):
#         # Передаём текущего пользователя как instance в форму
#         if form_class is None:
#             form_class = self.get_form_class()
#         return form_class(instance=self.request.user, **self.get_form_kwargs())

#     def form_valid(self, form):
#         # Сохраняем изменения
#         form.save()
#         return super().form_valid(form)

class UserProfileUpdateView(LoginRequiredMixin, FormView):
    title = "Редактирование профиля"
    template_name = "user_profile/update.html"
    form_class = UserProfileForm
    success_url = "/user_profile/"

    def get_form(self, form_class=None):
        # Передаём текущего пользователя в форму
        if form_class is None:
            form_class = self.get_form_class()
        return form_class(instance=self.request.user, **self.get_form_kwargs())

    def form_valid(self, form):
        # Сохраняем данные пользователя
        form.save()
        return super().form_valid(form)

    def form_invalid(self, form):
        # Для отладки
        print("Ошибки формы:", form.errors)
        return super().form_invalid(form)
    
# **********************************************************************************************
# Судебный приказ
# **********************************************************************************************
class CourtOrderListView(LoginRequiredMixin, FilterView):
    title = "Судебный приказ"
    model = CourtOrder
    template_name = "court_order/list.html"
    paginate_by = 20

    def get_queryset(self):
        queryset = super().get_queryset()
        # queryset = queryset.filter(claimant_id=self.request.user.id defendant_id=self.request.user.id)
        queryset = queryset.filter(defendant_id=self.request.user.id)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Pass the filterset to the template - it provides the form.
        # context['filterset'] = self.filterset
        return context


class CourtOrderDetailView(LoginRequiredMixin, DetailView):
    title = "Просмотр"
    model = CourtOrder
    template_name = "court_order/detail.html"


# **********************************************************************************************
# Заявление на отмену
# **********************************************************************************************


class CancellationRequestListView(LoginRequiredMixin, FilterView):
    title = "Заявления об отмене судебных приказов"
    model = CancellationRequest
    template_name = "cancellation-request/list.html"
    paginate_by = 10
    # filterset_class = RequestFilter
    # context_object_name = 'communities'

    def get_queryset(self):
        queryset = super().get_queryset()
        queryset = queryset.filter(user_id=self.request.user.id)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        return context


class CancellationRequestDetailView(LoginRequiredMixin, DetailView):
    title = "Просмотр"
    model = CancellationRequest
    template_name = "cancellation-request/detail.html"

    def get_object(self, queryset=None):
        obj = super().get_object(queryset)
        get_cancellation_request(self.request, obj.id)
        return obj

    def get_context_data(self, *args, **kwargs):
        context = super(CancellationRequestDetailView, self).get_context_data(*args, **kwargs)
        # object = context['object']
        # context['stat'] = RequestStat.objects.filter(document_id=object.id).all()
        return context


class CancellationRequestCreateView(LoginRequiredMixin, CreateView):
    title = "Добавление"
    model = CancellationRequest
    template_name = "cancellation-request/create.html"
    success_url = reverse_lazy("cancellation-request-detail")
    fields = ["content", "status"]

    def get_court_order(self):
        return get_court_order(self.request, self.kwargs["court_order_id"])

    def get_initial(self, *args, **kwargs):
        initial = super(CancellationRequestCreateView, self).get_initial(**kwargs)
        court_order = self.get_court_order()
        initial["court_order"] = court_order.id
        initial["user"] = court_order.defendant
        initial[
            "content"
        ] = """Я не согласна с судебным приказом и с требованиями, изложенными взыскателем в заявлении о выдаче судебного приказа,
поскольку представленная в нем сумма задолженности не соответствует действительности.
На основании изложенного, руководствуясь статьями 128, 129 Гражданского процессуального кодекса РФ,

ПРОШУ:
Отменить судебный приказ"""
        return initial

    def form_valid(self, form):
        court_order = self.get_court_order()
        form.instance.court_order_id = court_order.id
        form.instance.user = court_order.defendant
        form.instance.registration_date = datetime.now()
        obj = form.save()
        self.success_url = reverse("cancellation-request-detail", kwargs={"pk": obj.id})
        return super(CancellationRequestCreateView, self).form_valid(form)

    def get_context_data(self, *args, **kwargs):
        context = super(CancellationRequestCreateView, self).get_context_data(*args, **kwargs)
        return context


class CancellationRequestUpdateView(LoginRequiredMixin, UpdateView):
    title = "Редактирование"
    model = CancellationRequest
    fields = ["content", "status"]
    template_name = "cancellation-request/update.html"

    def get_success_url(self):
        pk = self.kwargs["pk"]
        return reverse("cancellation-request-detail", kwargs={"pk": pk})


class CancellationRequestDeleteView(LoginRequiredMixin, DeleteView):
    title = "Удаление"
    model = CancellationRequest
    template_name = "cancellation-request/delete.html"
    success_url = reverse_lazy("cancellation-request-list")


# **********************************************************************************************
# Шаблон
# **********************************************************************************************
class TemplateListView(LoginRequiredMixin, FilterView):
    title = "Шаблоны"
    model = Template
    template_name = "template/list.html"
    paginate_by = 20
    filterset_class = TemplateFilter

    def get_queryset(self):
        queryset = super().get_queryset()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        return context


# class TemplateDetailView(LoginRequiredMixin, DetailView):
#     title = 'Просмотр'
#     model = Template
#     template_name = 'template/detail.html'


class TemplateDetailView(LoginRequiredMixin, DetailView):
    title = "Просмотр"
    model = Template
    template_name = "template/detail.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        template_object = self.object

        # Читаем содержимое .docx
        preview_content = []
        if template_object.file:
            try:
                doc_path = template_object.file.path
                doc = Document(doc_path)

                # Преобразуем содержимое в список строк
                for paragraph in doc.paragraphs:
                    preview_content.append(escape(paragraph.text))  # Экранируем текст
            except Exception as e:
                preview_content = [f"Ошибка при загрузке содержимого: {escape(str(e))}"]

        context["preview_content"] = preview_content
        return context
